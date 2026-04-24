import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

labs_info = {
    "Lab 5: Input Controls (Spinners & Pickers)": {
        "Q1": "Develop an application named \"Vehicle Parking Registration\" where the user can register their vehicle for parking. The app should include a spinner that allows users to select the type of vehicle (e.g., car, bike, etc.) and text fields for entering the vehicle number and RC number. Upon clicking the submit button, the entered details should be displayed in a separate view, providing the user with options to either confirm the details or edit them. Once the user confirms the information, a toast message should appear showing a unique serial number to confirm the parking allotment.",
        "Q2": "Design and develop a \"Travel Ticket Booking\" app where users can book tickets by selecting the source and destination from dropdown lists (spinners) and choosing the date of travel using a date picker. Include a toggle button to let users specify whether they want a one-way ticket or a round-trip ticket. The app should have \"Submit\" and \"Reset\" buttons. When the \"Submit\" button is clicked, display the entered details on a new screen in a structured format. The \"Reset\" button should clear all input fields and set the date picker to the current system date and time.",
        "Q3": "Design and develop a \"Movie Ticket Booking\" app where users can book tickets by selecting the movie and theatre from dropdown menus (spinners), the date of the show using a date picker, and the preferred showtime using a time picker. Include a toggle button to let users choose between a standard ticket or a premium ticket. If the \"premium\" option is selected, ensure the \"Submit\" button becomes clickable only after 12:00 PM. The app should have \"Book Now\" and \"Reset\" buttons. When the user clicks \"Book Now,\" display the entered details along with available seats for the selected show in a new screen in a well-organized format. The \"Reset\" button should clear all fields and reset the date picker to the current system date. Validate all inputs to ensure correct data entry."
    },
    "Lab 6: Introduction to Menu": {
        "Q1": "Design a home page for the \"XYZ Fitness Center\" using an Options Menu with the following Requirements. a. Create a simple option menus where in once you click the “menu options”, the option items should get displayed which are \"Workout Plans,\" \"Trainers,\" \"Membership\". \"Workout Plans\" should display the list of workout programs (e.g., Weight Loss, Cardio). \"Trainers\" should display the names and specializations of trainers with their photos. \"Membership\" should show the membership packages with pricing details. b. This option menu uses images/icons instead of textual content. This textual content should represent “Contact US”, “About Us”,” Homepage”. Once the user clicks on the corresponding icons it should display corresponding content."
    },
    "Lab 7: Context and Pop Up Menu": {
        "Q1": "Design and develop an application that displays a list of installed applications on your device. On long press of any application, display the following options: a) Show whether the application is a system app or user-installed. b) Provide options to open the app, uninstall it, or view its details (e.g., version, storage usage). c) Indicate whether the app has special permissions enabled, like location or camera access. When an option like \"View Details\" is selected, navigate to a detailed view showing the app's permissions, size, and version. For the \"Uninstall\" option, prompt the user for confirmation before proceeding.",
        "Q2": "Create a View with the name “My Menu ” which contains an image as an icon. On click of that icon it shows a submenu as ”Image -1”,”Image -2” .On click of each item that particular image should be displayed with the corresponding content in the Toast.",
        "Q3": "Create a view that displays textual content providing a description of \"Digital Transformation\". Add a filter option with submenus: a. \"Search Keywords\" - Allow the user to input keywords that can be searched within the content. b. \"Highlight\" - Highlight specific words or phrases provided by the user in the document. c. \"Sort\" - Provide an option to sort the entire content either alphabetically or by relevance to the search keywords."
    },
    "Lab 8: Android SQLite & Shared Preferences": {
        "Q1": "Write a program to create a \"Task Manager\" application with the following features: a) A form to create and save tasks with fields like task name, due date, and priority level (High, Medium, Low). b) A list view to display all the saved tasks with their details. c) An option to edit or delete any task from the list.",
        "Q2": "Write a program to add grocery items along with its cost into the database and display the total cost of all the item selected by the user. Items has to be displayed through spinner.",
        "Q3": "Create an application “Movie Review” to create a movie review to do the following: a) User can write a movie review by stating movie name, year,giving points ranging from 1-5 and save details in a database. b) User can view the movie review for which review are already defined .Movie names has to be displayed using listview,and the selected movie’s details should be displayed in a table.",
        "Q4": "Make use of SQLite browser to view the database ,tables created. Modify the tables and upload the same into your project and view the modified data through your app.",
        "Q5": "Demonstrates the use of the Shared Preferences through an android application which displays a screen with some text fields. Save the value when the application is closed and brought back when it is opened again."
    }
}

def clean_code(content, is_xml):
    if is_xml:
        allowed_attrs = {'android:id', 'android:text', 'android:hint', 'android:orientation', 'android:inputType'}
        
        def repl_tag(match):
            tag = match.group(0)
            if tag.startswith('</'):
                return tag
            if tag.startswith('<!--'):
                return ''
            if tag.startswith('<?'):
                return tag
                
            tag = re.sub(r'\s+', ' ', tag)
            
            m = re.match(r'^<([^\s>]+)', tag)
            if not m:
                return tag
            tag_name = m.group(1)
            
            attrs = re.findall(r'([a-zA-Z0-9_\-:]+)=([\'"][^\'"]*[\'"])', tag)
            kept_attrs = []
            for name, val in attrs:
                if name in allowed_attrs:
                    kept_attrs.append(f"{name}={val}")
                    
            is_self_closing = tag.endswith('/>') or tag.endswith('/ >')
            
            res = f"<{tag_name}"
            if kept_attrs:
                res += " " + " ".join(kept_attrs)
                
            if is_self_closing:
                res += " />"
            else:
                res += ">"
                
            return res
            
        processed_xml = re.sub(r'<[^>]+>', repl_tag, content)
        
        processed_xml = re.sub(r'>\s*<', '>\n<', processed_xml)
        lines = [L.strip() for L in processed_xml.split('\n') if L.strip()]
        
        formatted_lines = []
        indent_level = 0
        for line in lines:
            if line.startswith('</'):
                indent_level = max(0, indent_level - 1)
                formatted_lines.append('    ' * indent_level + line)
            elif line.endswith('/>') or line.startswith('<?'):
                formatted_lines.append('    ' * indent_level + line)
            else:
                formatted_lines.append('    ' * indent_level + line)
                if not line.startswith('<!--'):
                    indent_level += 1
                    
        return '\n'.join(formatted_lines)
    else:
        # Clean Java/Kotlin code
        content = re.sub(r'^package com\.example\..*?;\s*', '', content, flags=re.MULTILINE)
        content = re.sub(r'^import .*?;\s*', '', content, flags=re.MULTILINE)
        content = re.sub(r'\n\s*\n', '\n\n', content)
        return content.strip()

def collect_files(base_path):
    # Only look inside app/src/main
    main_path = os.path.join(base_path, 'app', 'src', 'main')
    files_collected = []
    
    if not os.path.exists(main_path):
        return files_collected
        
    exclude_files = {
        "ic_launcher_background.xml", "ic_launcher_foreground.xml", "ic_launcher.xml", "ic_launcher_round.xml",
        "backup_rules.xml", "data_extraction_rules.xml",
        "colors.xml", "strings.xml", "themes.xml",
        "AndroidManifest.xml"
    }
        
    for root, dirs, files in os.walk(main_path):
        for f in files:
            if f in exclude_files:
                continue
            if f.endswith('.java') or f.endswith('.xml'):
                files_collected.append(os.path.join(root, f))
    
    return files_collected

def main():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    base_dir = r"c:\Users\ratul\AndroidStudioProjects"
    
    for lab_name, questions in labs_info.items():
        # Extrapolate lab folder name from "Lab X:"
        lab_match = re.search(r'Lab (\d+)', lab_name)
        if not lab_match:
            continue
        lab_folder = f"LAB{lab_match.group(1)}"
        lab_path = os.path.join(base_dir, lab_folder)
        
        doc.add_heading(lab_name, level=1)
        
        for q_key, q_text in questions.items():
            q_dir = os.path.join(lab_path, q_key)
            if not os.path.exists(q_dir):
                if lab_folder == "LAB6" and q_key == "Q1" and os.path.exists(os.path.join(lab_path, "app")):
                    q_dir = lab_path
                else:
                    print(f"Skipping {q_dir} (not found)")
                    continue
                
            q_paragraph = doc.add_paragraph()
            q_run = q_paragraph.add_run(f"{q_key}: {q_text}")
            q_run.bold = True
            
            files = collect_files(q_dir)
            for file_path in files:
                file_name = os.path.basename(file_path)
                
                doc.add_heading(file_name, level=2)
                
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                    
                cleaned_content = clean_code(content, file_name.endswith('.xml'))
                
                # Add code block
                code_para = doc.add_paragraph()
                code_run = code_para.add_run(cleaned_content)
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(9)
                
    output_path = os.path.join(base_dir, "Final_Lab_Record_v3.docx")
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == '__main__':
    main()
